from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


root = Path('tmp/material-map-validation/source')
root.mkdir(parents=True, exist_ok=True)

(root / 'project-brief.md').write_text(
    '# Product research brief\n\nOffline-first material mapping is the launch priority.\n', encoding='utf-8'
)
(root / 'meeting-notes.txt').write_text(
    '2026-07-30\n\nTeam agreed to validate imports, relationships, and AI analysis.\n', encoding='utf-8'
)
(root / 'research-tracker.csv').write_text(
    'item,owner,status\nImport validation,Lin,complete\nModel analysis,Chen,planned\n', encoding='utf-8'
)
(root / 'reference-page.html').write_text(
    '<html><head><title>Reference page</title></head><body><h1>Workflow reference</h1><p>Materials connect through evidence.</p></body></html>', encoding='utf-8'
)

document = Document()
document.add_heading('Validation report', level=1)
document.add_paragraph('This DOCX verifies office-document extraction in the local workspace.')
document.save(root / 'validation-report.docx')

pdf = canvas.Canvas(str(root / 'timeline.pdf'), pagesize=A4)
pdf.setFont('Helvetica', 14)
pdf.drawString(72, 760, 'Validation timeline')
pdf.setFont('Helvetica', 11)
pdf.drawString(72, 730, 'The local importer preserves files and extracts material content.')
pdf.save()

image = Image.new('RGB', (1200, 360), 'white')
draw = ImageDraw.Draw(image)
draw.text((48, 70), 'OCR validation: material map evidence', fill='black', font_size=36)
image.save(root / 'ocr-sample.png')

print(root.resolve())
